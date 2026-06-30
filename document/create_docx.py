from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_doc():
    doc = Document()
    
    # Configure default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # Title
    title = doc.add_paragraph('TỔNG QUAN TÀI LIỆU')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.bold = True
    title.runs[0].font.size = Pt(16)
    
    subtitle = doc.add_paragraph('Nghiên cứu Ứng dụng Mô hình NeuralProphet, ARIMA và LSTM trong Bài toán Dự báo Luồng Không lưu')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.bold = True

    # 1. Giới thiệu tổng quan
    doc.add_heading('1. Giới thiệu tổng quan', level=1)
    doc.add_heading('1.1. Bối cảnh nghiên cứu', level=2)
    doc.add_paragraph('Quản lý luồng không lưu (ATFM) đóng vai trò then chốt trong việc bảo đảm sự cân bằng giữa nhu cầu khai thác và năng lực của hệ thống quản lý không lưu. Trong đó, công tác dự báo chuỗi thời gian là một thành phần cốt lõi, cho phép Kiểm soát viên không lưu (KSVKL) lường trước lưu lượng tàu bay, từ đó đưa ra các quyết định chiến lược về điều phối sức chứa phân khu và phân bổ nguồn lực. Việc lựa chọn một mô hình dự báo tối ưu có tác động trực tiếp đến độ chính xác của dự báo, hiệu năng tính toán và tính minh bạch trong việc ra quyết định điều hành.')
    doc.add_paragraph('Nhận thức được tầm quan trọng đó, tài liệu này tiến hành đánh giá ba phương pháp dự báo chuỗi thời gian tiêu biểu: NeuralProphet, ARIMA và LSTM. Nghiên cứu tập trung cung cấp các số liệu đánh giá (benchmark) đã được xác minh từ các nguồn tài liệu bình duyệt (peer-reviewed), nhằm làm cơ sở khoa học để đề xuất giải pháp cho hệ thống AeroCast VAA.')

    doc.add_heading('1.2. Mục tiêu nghiên cứu', level=2)
    doc.add_paragraph('Tài liệu này tập trung giải quyết các mục tiêu cụ thể sau:')
    doc.add_paragraph('1. Đánh giá đặc tính hiệu suất của mô hình NeuralProphet so với các phương pháp thống kê truyền thống.', style='List Number')
    doc.add_paragraph('2. Xác định các điều kiện khai thác phù hợp để ứng dụng mô hình ARIMA so với các phương pháp Học sâu (Deep Learning).', style='List Number')
    doc.add_paragraph('3. Phân tích những hạn chế của mô hình LSTM khi áp dụng vào các tập dữ liệu có quy mô nhỏ.', style='List Number')
    doc.add_paragraph('4. Đề xuất mô hình dự báo khả thi và hiệu quả nhất cho công tác quản lý luồng không lưu ngắn hạn.', style='List Number')

    # 2. Phương pháp
    doc.add_heading('2. Phương pháp Nghiên cứu', level=1)
    doc.add_heading('2.1. Nguồn thu thập dữ liệu', level=2)
    doc.add_paragraph('Các cơ sở dữ liệu học thuật và kho lưu trữ sau đây đã được sử dụng làm cơ sở tham khảo:')
    doc.add_paragraph('arXiv: Hệ thống lưu trữ tài liệu khoa học mở trong lĩnh vực khoa học máy tính và thống kê.', style='List Bullet')
    doc.add_paragraph('IEEE Xplore: Thư viện số chuyên ngành kỹ thuật và công nghệ điện tử.', style='List Bullet')
    doc.add_paragraph('Google Scholar: Nền tảng tra cứu tài liệu học thuật tổng hợp.', style='List Bullet')
    doc.add_paragraph('Kaggle Competitions: Nguồn cung cấp dữ liệu thực nghiệm từ các cuộc thi dự báo M4 và M5.', style='List Bullet')

    doc.add_heading('2.2. Các chỉ số đánh giá hiệu năng (Evaluation Metrics)', level=2)
    doc.add_paragraph('Nghiên cứu ưu tiên sử dụng ba chỉ số đánh giá sai số cốt lõi. Mỗi chỉ số được lựa chọn đều mang ý nghĩa quan trọng trong thực tiễn điều hành bay:')

    doc.add_heading('2.2.1. Chỉ số RMSE (Root Mean Square Error)', level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('RMSE = √[ (1/n) * Σ(yi - ŷi)² ]').italic = True
    p2 = doc.add_paragraph()
    p2.add_run('Ứng dụng trong thực tiễn kiểm soát không lưu: ').bold = True
    p2.add_run('Đặc tính của RMSE là bình phương sai số trước khi tính trung bình, qua đó khuếch đại và làm nổi bật các sai lệch lớn. Trong thực tiễn khai thác, nếu mô hình dự báo sai lệch một tàu bay trong nhiều chu kỳ liên tiếp, KSVKL vẫn có khả năng điều hòa luồng không lưu một cách an toàn. Tuy nhiên, nếu hệ thống dự báo sai lệch một số lượng lớn tàu bay chỉ trong một chu kỳ duy nhất, phân khu kiểm soát có nguy cơ rơi vào trạng thái quá tải (Sector Overload), uy hiếp trực tiếp đến an toàn bay. Do đó, RMSE là chỉ số đánh giá độ tin cậy cốt lõi, đảm bảo mô hình hạn chế tối đa các sai lệch mang tính rủi ro cao.')

    doc.add_heading('2.2.2. Chỉ số MAPE (Mean Absolute Percentage Error)', level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('MAPE = (100% / n) * Σ| (yi - ŷi) / yi |').italic = True
    p2 = doc.add_paragraph()
    p2.add_run('Ứng dụng trong thực tiễn kiểm soát không lưu: ').bold = True
    p2.add_run('Chỉ số MAPE biểu diễn sai số dưới định dạng phần trăm (%), hỗ trợ KSVKL đánh giá trực quan mức độ nghiêm trọng của sự sai lệch so với tổng lưu lượng khai thác thực tế. Việc dự báo chênh lệch 2 tàu bay trên tổng lưu lượng 40 tàu bay (sai số 5%) nằm trong ngưỡng an toàn cho phép. Tuy nhiên, nếu lưu lượng thực tế chỉ có 3 tàu bay nhưng hệ thống sai lệch 2 tàu bay (sai số 66.6%), điều này phản ánh sự bất ổn định của mô hình tính toán. MAPE giúp chuẩn hóa việc đánh giá mức độ rủi ro bất chấp quy mô năng lực của từng phân khu.')

    doc.add_heading('2.2.3. Chỉ số MASE (Mean Absolute Scaled Error)', level=3)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('MASE = Tử số (Sai số mô hình AI) / Mẫu số (Sai số Dự báo Naive)').italic = True
    p2 = doc.add_paragraph()
    p2.add_run('Ứng dụng trong thực tiễn kiểm soát không lưu: ').bold = True
    p2.add_run('Chỉ số MASE thực hiện việc so sánh sai số của mô hình đề xuất (tử số) với sai số của phương pháp dự báo cơ sở (Dự báo Naive - mẫu số). Trong ngữ cảnh quản lý không lưu, dự báo Naive tương đương với việc ngoại suy trực tiếp lưu lượng hiện tại cho chu kỳ tiếp theo mà không thông qua phân tích xu hướng.')
    doc.add_paragraph('Trường hợp MASE > 1: Hiệu suất của mô hình hệ thống thấp hơn phương pháp ngoại suy cơ bản, chứng tỏ thuật toán không mang lại giá trị cải thiện trong thực tiễn điều hành.', style='List Bullet')
    doc.add_paragraph('Trường hợp MASE < 1: Khẳng định sự ưu việt của mô hình trong việc phân tích và dự báo chính xác động học không lưu. Đây là thước đo khách quan nhằm đánh giá tính khả thi và độ tin cậy của hệ thống.', style='List Bullet')

    # 3. NeuralProphet
    doc.add_heading('3. Cơ sở Lý thuyết Mô hình NeuralProphet', level=1)
    doc.add_heading('3.1. Kiến trúc mô hình và Ứng dụng thực tiễn', level=2)
    doc.add_paragraph('Mô hình NeuralProphet tích hợp các thành phần phân tích thống kê cổ điển kết hợp cùng kiến trúc mạng nơ-ron thông qua phương trình:')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('ŷ(t) = Trend(t) + Seasonality(t) + AR-Net(t) + Regressors(t)').italic = True
    
    p2 = doc.add_paragraph()
    p2.add_run('Chi tiết các thành phần trong công tác Quản lý luồng không lưu:').bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('ŷ(t): ').bold = True
    p.add_run('Dự báo tổng lưu lượng tàu bay tiến nhập vào phân khu kiểm soát tại thời điểm t định trước.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Trend(t): ').bold = True
    p.add_run('Xu hướng thay đổi dài hạn của luồng không lưu. Trong thực tiễn khai thác, yếu tố này phản ánh sự tăng trưởng lưu lượng trong các giai đoạn cao điểm du lịch hoặc tiến trình phục hồi mạng đường bay.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Seasonality(t): ').bold = True
    p.add_run('Yếu tố chu kỳ. Hỗ trợ hệ thống nhận diện các khung giờ cao điểm cố định trong ngày (ví dụ: khung giờ cất cánh buổi sáng) hoặc sự phân bổ lưu lượng không đồng đều giữa các ngày trong tuần.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('AR-Net(t): ').bold = True
    p.add_run('Mạng nơ-ron tự hồi quy. Đây là cơ chế cốt lõi giúp mô hình tiếp thu xu hướng động học của dòng không lưu. Dựa trên số lượng tàu bay vừa vượt qua các điểm báo cáo bắt buộc (Ví dụ: điểm ANLOC), mô hình tự động suy diễn và phân bổ chính xác số lượng tàu bay chuẩn bị tiến nhập vào phân khu tiếp cận trong các chu kỳ kế tiếp.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Regressors(t): ').bold = True
    p.add_run('Biến ngoại sinh. Yếu tố này đóng vai trò quyết định trong việc hiệu chỉnh kết quả dự báo dưới tác động của môi trường: Điều kiện khí tượng tiêu chuẩn (Regressors tĩnh) và Điều kiện khí tượng bất lợi (mây dông). Khi đó hệ thống tự động tính toán và điều chỉnh giảm giá trị ŷ(t), bảo đảm sản lượng dự báo đồng bộ với năng lực tiếp thu thực tế của đường băng, ngăn ngừa sự cố tắc nghẽn cục bộ.')

    # 4. ARIMA
    doc.add_heading('4. Cơ sở Lý thuyết Mô hình ARIMA', level=1)
    doc.add_heading('4.1. Cơ sở Toán học và Ứng dụng Khai thác', level=2)
    doc.add_paragraph('Phương trình tổng quát của mô hình ARIMA(p, d, q) được biểu diễn như sau:')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Φ(B)(1-B)ᵈ y_t = Θ(B)ε_t').italic = True
    
    p2 = doc.add_paragraph()
    p2.add_run('Ứng dụng các tham số trong công tác điều hành bay:').bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('y_t: ').bold = True
    p.add_run('Số lượng tàu bay thực tế được ghi nhận tại phân đoạn thời gian t.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('B: ').bold = True
    p.add_run('Toán tử dịch lùi (Backshift), đại diện cho các giá trị lưu lượng trong các chu kỳ khai thác trước đó.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('d: ').bold = True
    p.add_run('Bậc vi phân (Differencing). Thành phần này có chức năng loại bỏ tính không dừng (non-stationary) của chuỗi dữ liệu, đóng vai trò như một cơ chế làm phẳng sự biến động đột ngột của luồng không lưu.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Φ(B): ').bold = True
    p.add_run('Đa thức Tự hồi quy (AR). Phân tích sự phụ thuộc tuyến tính của lưu lượng hiện tại vào dữ liệu quá khứ. Trong thực tiễn, nếu lưu lượng tàu bay tăng trưởng liên tục qua các chu kỳ, thành phần AR sẽ thiết lập cơ sở dự báo duy trì đà tăng trong chu kỳ tiếp theo.')
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Θ(B): ').bold = True
    p.add_run('Đa thức Trung bình động (MA). Có chức năng đo lường và triệt tiêu các sai số ngẫu nhiên. Ví dụ trường hợp chuyến bay ưu tiên xin cấp huấn lệnh hạ cánh đột xuất gây nhiễu loạn biểu đồ, thành phần MA sẽ tự động hiệu chỉnh và hấp thụ sai số này.')

    # 5. LSTM
    doc.add_heading('5. Cơ sở Lý thuyết Mô hình LSTM', level=1)
    doc.add_heading('5.1. Kiến trúc Mạng và Cơ chế Hoạt động trong Thực tiễn', level=2)
    doc.add_paragraph('Cấu trúc cốt lõi của một khối LSTM bao gồm Trạng thái ô (Cell State) hoạt động như một luồng bộ nhớ chủ đạo, được kiểm soát bởi ba cơ chế cổng (Gates):')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('f_t = σ(W_f · [h_{t-1}, x_t] + b_f)   (Cổng Quên)').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('i_t = σ(W_i · [h_{t-1}, x_t] + b_i)   (Cổng Nhập)').italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('o_t = σ(W_o · [h_{t-1}, x_t] + b_o)   (Cổng Xuất)').italic = True
    
    p2 = doc.add_paragraph()
    p2.add_run('Ứng dụng cơ chế vận hành trong quản lý không lưu:').bold = True
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Cổng quên (f_t): ').bold = True
    p.add_run('Trong quá trình điều hành bay, khi một tàu bay đã hạ cánh an toàn và thoát ly khỏi đường cất hạ cánh, thuật toán sẽ tính toán đưa giá trị f_t về 0, qua đó triệt tiêu dữ liệu của tàu bay đó khỏi bộ nhớ hệ thống nhằm giải phóng dung lượng phục vụ cho hoạt động giám sát các mục tiêu mới.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Cổng nhập (i_t): ').bold = True
    p.add_run('Khi hệ thống giám sát tự động (ADS-B/Radar) phát hiện một tín hiệu mục tiêu xuất hiện tại ranh giới Vùng thông báo bay (FIR), cổng nhập sẽ đánh giá các tham số vận tốc và phương vị. Nếu mục tiêu có quỹ đạo tiến nhập vào phân khu kiểm soát, dữ liệu sẽ lập tức được cập nhật.')
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run('Cổng xuất (o_t): ').bold = True
    p.add_run('Trích xuất dữ liệu tổng hợp để cấu thành kết quả dự báo. Căn cứ vào trạng thái tàu bay vùng tiếp cận, hệ thống sẽ xử lý và truyền tải số liệu dự báo lưu lượng lên Màn hình hiển thị của KSVKL (CWP), đóng vai trò thông tin tham khảo chiến lược.')

    # Save doc
    doc.save(r'D:\Project\AeroCast-VAA-System-V2\AeroCast-VAA-System-V2\document\Literature_Review_HocThuat.docx')

create_doc()
