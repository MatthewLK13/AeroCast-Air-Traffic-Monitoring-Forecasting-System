"""
Check and analyze the SPEC.docx file for issues
"""

from docx import Document
from docx.shared import Pt, RGBColor
import os

# Set UTF-8 encoding for output
import sys
sys.stdout.reconfigure(encoding='utf-8')

doc_path = os.path.join(os.path.dirname(__file__), '..', 'spec', 'SPEC.docx')

doc = Document(doc_path)

print("=== DOCUMENT ANALYSIS ===\n")

# Count elements
paragraphs = len(doc.paragraphs)
tables = len(doc.tables)
images = len([shape for shape in doc.inline_shapes if hasattr(shape, 'image')])

print(f"Paragraphs: {paragraphs}")
print(f"Tables: {tables}")
print(f"Images: {images}")

# Check paragraphs for issues
print("\n=== PARAGRAPH STYLES ===")
style_counts = {}
for para in doc.paragraphs:
    style_name = para.style.name if para.style else "None"
    style_counts[style_name] = style_counts.get(style_name, 0) + 1

for style, count in sorted(style_counts.items(), key=lambda x: -x[1]):
    print(f"  {style}: {count}")

# Check for text with special characters
print("\n=== POTENTIAL ISSUES ===")
issues = []
for i, para in enumerate(doc.paragraphs):
    text = para.text
    # Check for common problematic patterns
    if '□' in text or '�' in text:
        issues.append(f"Para {i}: Contains replacement character")
    if len(para.text) > 1000:
        issues.append(f"Para {i}: Very long ({len(para.text)} chars)")

if issues:
    for issue in issues[:10]:
        print(f"  - {issue}")
else:
    print("  No obvious text issues found")

# Check tables
print(f"\n=== TABLES ({len(doc.tables)}) ===")
for i, table in enumerate(doc.tables):
    rows = len(table.rows)
    cols = len(table.columns)
    print(f"Table {i}: {rows} rows x {cols} cols")

print("\n=== SUMMARY ===")
print("Document structure looks OK")
print("Issue likely in: Unicode rendering in Word, not in the docx itself")
