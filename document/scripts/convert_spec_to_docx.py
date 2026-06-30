"""
Convert SPEC.md to SPEC.docx - Manual Construction
Builds the document manually to ensure proper formatting
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# Paths
BASE_DIR = os.path.dirname(os.path.dirname(__file__))
SPEC_DOCX = os.path.join(BASE_DIR, 'spec', 'SPEC.docx')
FIGURES_DIR = os.path.join(BASE_DIR, 'figures')

def add_title(doc, text, size=32, bold=True, color=None):
    """Add a styled title"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def add_heading_custom(doc, text, level=1):
    """Add heading"""
    p = doc.add_heading(text, level=level)
    return p

def add_body(doc, text):
    """Add body paragraph"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0):
    """Add bullet point"""
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def add_image_para(doc, img_name, width=Inches(6)):
    """Add centered image"""
    img_path = os.path.join(FIGURES_DIR, img_name)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=width)
        return True
    else:
        p = doc.add_paragraph(f"[Image not found: {img_name}]")
        return False

def add_code_block(doc, lines):
    """Add code block with dark background"""
    p = doc.add_paragraph()
    for line in lines:
        run = p.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)
    return p

def add_spacer(doc, size=0.2):
    """Add spacing paragraph"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(size * 12)
    return p

def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Set dark blue background
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        solidFill = OxmlElement('w:solidFill')
        solidFill.set(qn('w:val'), '1565C0')
        tcPr.append(solidFill)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, cell_text in enumerate(row_data):
            row_cells[i].text = str(cell_text)
            for paragraph in row_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    return table

def create_document():
    """Create the full document"""
    doc = Document()

    # ==================== TITLE PAGE ====================
    add_spacer(doc, 1)

    p = add_title(doc, 'AEROCAST VAA SYSTEM', 40, True, RGBColor(0x1F, 0x4E, 0x79))
    add_spacer(doc, 0.5)

    p = add_title(doc, 'TÀI LIỆU KỸ THUẬT MÔ HÌNH DỰ BÁO', 20, False)
    add_spacer(doc, 0.5)

    p = add_title(doc, 'Phiên bản: 2.0 (Mở rộng)', 12, False)
    p = add_title(doc, 'Ngày: 29/05/2026', 12, False)

    add_spacer(doc, 1)

    # Add system overview diagram
    add_image_para(doc, 'diagram_01_system_overview.png', Inches(7))
    p = doc.add_paragraph('Hình 1: Tổng quan hệ thống AeroCast')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    doc.add_page_break()

    # ==================== TABLE OF CONTENTS ====================
    add_heading_custom(doc, 'MỤC LỤC', 1)

    toc_items = [
        '1. Tổng quan hệ thống',
        '2. Bài toán dự báo là gì?',
        '3. Tại sao không dùng công thức toán đơn giản?',
        '4. Giải thích Neural Network cho người không chuyên',
        '5. NeuralProphet hoạt động như thế nào?',
        '6. Từng dòng code có ý nghĩa gì?',
        '7. Tại sao chọn các con số cụ thể này?',
        '8. So sánh các phương pháp dự báo',
        '9. Tại sao không dùng phương pháp khác?',
        '10. Các thuật ngữ và giải thích',
        '11. Kết luận và câu hỏi thường gặp',
    ]

    for item in toc_items:
        add_bullet(doc, item)

    doc.add_page_break()

    # ==================== PHAN 1: SYSTEM OVERVIEW ====================
    add_heading_custom(doc, 'PHẦN 1: TỔNG QUAN HỆ THỐNG', 1)

    add_heading_custom(doc, '1.1 Hệ thống AeroCast làm gì?', 2)

    add_body(doc, 'AeroCast là hệ thống giám sát và dự báo không lưu hàng không cho Phân khu 2 vùng trời Việt Nam. Hệ thống sử dụng:')

    add_bullet(doc, 'FlightRadar24 API - Cung cấp dữ liệu vị trí máy bay thực tế')
    add_bullet(doc, 'Open-Meteo API - Cung cấp dữ liệu thời tiết (mưa, gió, mây, tầm nhìn)')
    add_bullet(doc, 'SQLite Database - Lưu trữ dữ liệu lịch sử')
    add_bullet(doc, 'NeuralProphet Model - Dự báo số máy bay trong 30 phút tới')

    add_spacer(doc, 0.3)

    add_body(doc, 'Mục tiêu: Giúp các controller điều phối không lưu hiệu quả hơn bằng cách dự đoán trước lưu lượng máy bay.')

    add_heading_custom(doc, '1.2 Chúng ta đang có dự đoán gì?', 2)

    add_body(doc, 'Đầu vào (Input):')
    add_bullet(doc, 'Số máy bay tại 12 timestep trước đó (60 phút)')
    add_bullet(doc, 'Thời tiết hiện tại: lượng mưa, tốc độ gió, độ che phủ mây, tầm nhìn')
    add_bullet(doc, 'Buffer counts: số máy bay đang di chuyển vào sector từ 4 hướng')

    add_body(doc, 'Đầu ra (Output):')
    add_bullet(doc, 'Dự báo số máy bay trong 6 bước tiếp theo (30 phút)')
    add_bullet(doc, 'Khoảng dao động 80% confidence [y_10, y_90]')

    add_spacer(doc, 0.3)

    # ==================== PHAN 2: WHAT IS FORECASTING ====================
    add_heading_custom(doc, 'PHẦN 2: BÀI TOÁN DỰ BÁO LÀ GÌ?', 1)

    add_heading_custom(doc, '2.1 Dự báo - Khái niệm đời thường', 2)

    add_body(doc, 'Dự báo = Dự đoán tương lai dựa trên quá khứ. Ví dụ:')
    add_bullet(doc, 'Dự báo thời tiết: Thấy mây đen kéo đến -> đoán sắp mưa')
    add_bullet(doc, 'Dự báo kẹt xe: 7h thứ Hai -> đường sẽ kẹt vì lượng kẹt')
    add_bullet(doc, 'Dự báo lượng khách: Sáng sớm và chiều muộn đông khách')

    add_heading_custom(doc, '2.2 Dự báo trong khoa học máy tính', 2)

    add_body(doc, 'Công thức đơn giản:')
    add_bullet(doc, 'y_t = số máy bay tại thời điểm t')
    add_bullet(doc, 'y_{t-1} = số máy bay 5 phút trước')
    add_bullet(doc, 'y_{t-2} = số máy bay 10 phút trước')
    add_body(doc, 'Mục tiêu: Tìm công thức tính y_{t+k} (k bước trong tương lai)')

    add_heading_custom(doc, '2.3 Tại sao dự báo khó?', 2)

    add_bullet(doc, 'QUÁ KHỨ KHÔNG HOÀN TOÀN LẶP LẠI')
    add_bullet(doc, '   Hôm qua 8h có 5 máy bay, hôm nay 8h có 8 máy bay')
    add_bullet(doc, 'CÓ NHIỀU YẾU TỐ ẢNH HƯỞNG')
    add_bullet(doc, '   Thời tiết: Mưa -> máy bay phải chờ')
    add_bullet(doc, '   Sự kiện: Thêm chuyến bay mới')
    add_bullet(doc, 'SỰ BẤT ĐỊNH (UNCERTAINTY)')
    add_bullet(doc, '   Không thể đoán chính xác bao nhiêu máy bay')
    add_bullet(doc, '   Chỉ có thể nói: "Có thể từ 4 đến 8 máy bay"')

    add_spacer(doc, 0.3)

    # ==================== PHAN 3: WHY NOT SIMPLE MATH ====================
    add_heading_custom(doc, 'PHẦN 3: TẠI SAO KHÔNG DÙNG CÔNG THỨC TOÁN ĐƠN GIẢN?', 1)

    add_heading_custom(doc, '3.1 Phương pháp trung bình không đủ', 2)

    add_body(doc, 'Phương pháp đơn giản nhất là lấy trung bình các giá trị trước đó. Tuy nhiên:')

    add_bullet(doc, 'Tất cả các ngày đều cho ra cùng một kết quả')
    add_bullet(doc, 'Không phân biệt được 8h sáng (đông) vs 3h chiều (vắng)')
    add_bullet(doc, 'Không quan tâm đến thời tiết hay các yếu tố khác')

    add_heading_custom(doc, '3.2 ARIMA - Hạn chế lớn nhất', 2)

    add_body(doc, 'ARIMA (Auto-Regressive Integrated Moving Average) là phương pháp phổ biến, nhưng có giới hạn:')

    add_bullet(doc, 'CHỈ học được QUAN HỆ TUYẾN TÍNH')
    add_bullet(doc, '   Ví dụ: y = 2x + 1 (luôn theo đường thẳng)')
    add_bullet(doc, '   Thực tế: Giờ cao điểm đông gấp 3 LẦN, không phải gấp đôi')
    add_bullet(doc, 'CHỈ dùng được với MỘT biến (univariate)')
    add_bullet(doc, '   Không dùng được thời tiết, buffer counts trực tiếp')
    add_bullet(doc, '   Muốn dùng weather -> phải dùng ARIMAX (phức tạp hơn nhiều)')
    add_bullet(doc, 'SEASONALITY phải định nghĩa thủ công')
    add_bullet(doc, '   Phải tự cấu hình pattern theo giờ, theo ngày')

    add_heading_custom(doc, '3.3 Tại sao cần Neural Network?', 2)

    add_body(doc, 'Thực tế air traffic không tuân theo đường thẳng:')

    add_bullet(doc, '8h-9h: Đông dần (5 -> 8 máy bay)')
    add_bullet(doc, '9h-10h: Đông đỉnh (8 -> 10 máy bay)')
    add_bullet(doc, '10h-11h: Giảm nhanh (10 -> 6 máy bay)')
    add_bullet(doc, '11h-12h: Ổn định (6 -> 5 máy bay)')

    add_body(doc, 'Đây là đường CONG phức tạp - Neural Network có thể học được!')

    add_spacer(doc, 0.3)

    # ==================== PHAN 4: NEURAL NETWORK BASICS ====================
    add_heading_custom(doc, 'PHẦN 4: GIẢI THÍCH NEURAL NETWORK CHO NGƯỜI KHÔNG CHUYÊN', 1)

    add_heading_custom(doc, '4.1 Neural Network như bộ não nhân tạo', 2)

    add_body(doc, 'Neural Network mô phỏng cách hoạt động của bộ não con người:')
    add_bullet(doc, 'Input (X): Dữ liệu đầu vào (số máy bay, thời tiết)')
    add_bullet(doc, 'Hidden Layers: Các tầng ẩn xử lý thông tin')
    add_bullet(doc, 'Output (Y): Kết quả dự báo')
    add_bullet(doc, 'Weights: Trọng số - độ quan trọng của mỗi kết nối')

    add_image_para(doc, 'diagram_03_neural_network.png', Inches(6.5))
    p = doc.add_paragraph('Hình 2: Kiến trúc Neural Network với 2 hidden layers [32, 32]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    add_heading_custom(doc, '4.2 Neural Network hoạt động như thế nào?', 2)

    add_body(doc, 'Bước 1: Khởi tạo weights ngẫu nhiên')
    add_body(doc, 'Bước 2: Tính sai số (Error) - so sánh dự báo với thực tế')
    add_body(doc, 'Bước 3: Điều chỉnh weights để sai số giảm (Backpropagation)')
    add_body(doc, 'Bước 4: Lặp lại nhiều epoch cho đến khi sai số đủ nhỏ')

    add_body(doc, '"Học" = Điều chỉnh trọng số để dự báo chính xác hơn')

    add_spacer(doc, 0.3)

    # ==================== PHAN 5: NEURALPROPHET HOW IT WORKS ====================
    add_heading_custom(doc, 'PHẦN 5: NEURALPROPHET HOẠT ĐỘNG NHƯ THẾ NÀO?', 1)

    add_image_para(doc, 'diagram_00_neuralprophet_architecture.png', Inches(7))
    p = doc.add_paragraph('Hình 3: NeuralProphet Architecture - Các thành phần chính của mô hình')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    add_heading_custom(doc, '5.1 NeuralProphet = Prophet + AR-Net', 2)

    add_body(doc, 'NeuralProphet là sự kết hợp giữa:')
    add_bullet(doc, 'Prophet (Facebook): Dự báo chuỗi thời gian với Trend + Seasonality')
    add_bullet(doc, 'AR-Net (Autoregressive Neural Network): Điểm khác biệt chính - nhìn vào quá khứ để dự đoán tương lai')

    add_heading_custom(doc, '5.2 Các thành phần chính', 2)

    add_body(doc, 'TREND: Xu hướng dài hạn (tăng/giảm theo thời gian)')
    add_body(doc, 'SEASONALITY: Pattern lặp lại (theo ngày, tuần, năm)')
    add_body(doc, 'AR-NET: Autoregressive Neural Network - điểm khác biệt chính với Prophet')
    add_body(doc, 'REGRESSORS: Biến ngoại sinh (weather, buffer counts)')
    add_body(doc, 'QUANTILES: Ước lượng uncertainty (khoảng dao động)')

    add_heading_custom(doc, '5.3 Điểm khác biệt chính với Prophet', 2)

    add_body(doc, 'Prophet: y(t) = trend(t) + seasonality(t) + holidays(t)')
    add_body(doc, '   -> KHÔNG có autoregressive component!')
    add_spacer(doc, 0.1)
    add_body(doc, 'NeuralProphet: y(t) = trend(t) + seasonality(t) + AR-Net(...) + regressors')
    add_body(doc, '   -> CÓ autoregressive component! Nhìn vào quá khứ để dự đoán!')

    add_spacer(doc, 0.3)

    # AR-NET Section
    add_heading_custom(doc, '5.4 AR-NET - Điểm khác biệt chính với Prophet', 2)

    add_body(doc, 'AR-NET sử dụng 12 giá trị quá khứ gần nhất:')
    add_body(doc, '   y(t-1), y(t-2), ..., y(t-12) -> dự báo y(t+1), ..., y(t+6)')

    add_image_para(doc, 'diagram_04_ar_net.png', Inches(6.5))
    p = doc.add_paragraph('Hình 4: AR-NET Architecture')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    # Quantiles Section
    add_heading_custom(doc, '5.5 Quantiles - Ước lượng sự bất định', 2)

    add_body(doc, 'QUANTILES giúp ước lượng uncertainty - "CHÚNG TÔI KHÔNG CHẮC CHẮN 100%"')

    add_body(doc, 'Thay vì dùng 1 model với MSE loss, NeuralProphet dùng 3 outputs:')
    add_bullet(doc, 'Model 1: dự báo 10th percentile (q=0.1) - 10% chance actual < prediction')
    add_bullet(doc, 'Model 2: dự báo MEDIAN (50th percentile) - 50% chance')
    add_bullet(doc, 'Model 3: dự báo 90th percentile (q=0.9) - 10% chance actual > prediction')

    add_body(doc, 'Kết quả: 80% CONFIDENCE INTERVAL [y_10, y_90]')

    add_image_para(doc, 'diagram_05_quantiles.png', Inches(6.5))
    p = doc.add_paragraph('Hình 5: Quantiles - Khoảng dao động 80% confidence')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    # Seasonality Section
    add_heading_custom(doc, '5.6 Seasonality - Tính chu kỳ', 2)

    add_body(doc, 'DAILY SEASONALITY: Pattern theo giờ trong ngày')
    add_bullet(doc, '6h-9h: Sáng sớm -> đông dần')
    add_bullet(doc, '14h-16h: Chiều -> đông đỉnh')
    add_bullet(doc, '3h-5h: Sáng sớm -> vắng')

    add_body(doc, 'WEEKLY SEASONALITY: Pattern theo ngày trong tuần')
    add_bullet(doc, 'Thứ 2 -> đông nhất (sau cuối tuần)')
    add_bullet(doc, 'Chủ nhật -> vắng nhất')

    add_image_para(doc, 'diagram_08_seasonality.png', Inches(6.5))
    p = doc.add_paragraph('Hình 6: Seasonality - Pattern theo ngày và tuần')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    add_spacer(doc, 0.3)

    # ==================== PHAN 6: CODE EXPLANATION ====================
    add_heading_custom(doc, 'PHẦN 6: TỪNG DÒNG CODE CÓ Ý NGHĨA GÌ?', 1)

    add_heading_custom(doc, '6.1 Cấu hình Model NeuralProphet', 2)

    code = [
        'm = NeuralProphet(',
        '    n_lags=12,           # Dùng 12 điểm dữ liệu QUÁ KHỨ (60 phút)',
        '    n_forecasts=6,       # Dự báo 6 bước TƯƠNG LAI (30 phút)',
        '    ar_layers=[32, 32],   # Neural network có 2 tầng ẩn, mỗi tầng 32 neuron',
        '    quantiles=[0.1, 0.9], # Khoảng dao động 10%-90% (80% confidence)',
        '    yearly_seasonality=False, # TẮT seasonality theo năm (data < 1 năm)',
        '    weekly_seasonality=True,  # BẬT seasonality theo tuần',
        '    daily_seasonality=True,   # BẬT seasonality theo ngày',
        ')'
    ]
    add_code_block(doc, code)

    add_heading_custom(doc, '6.2 Preprocessing Pipeline', 2)

    code = [
        '# Resample về 5 phút và interpolate',
        'df_prophet = df_prophet.set_index("ds")',
        '    .resample("5min").mean()',
        '    .interpolate(method="linear")',
        '    .bfill().ffill()',
        '',
        '# Thêm noise nhỏ để tránh zero variance',
        'df_prophet[col] = df_prophet[col].astype(float)',
        '    + np.random.normal(0, 1e-5, size=len(df_prophet))'
    ]
    add_code_block(doc, code)

    add_heading_custom(doc, '6.3 Training', 2)

    code = [
        '# Train model với minimal mode (tự chọn epochs phù hợp)',
        'm.fit(df_train, freq="5min", minimal=True)'
    ]
    add_code_block(doc, code)

    add_spacer(doc, 0.3)

    # ==================== PHAN 7: HYPERPARAMETERS ====================
    add_heading_custom(doc, 'PHẦN 7: TẠI SAO CHỌN CÁC CON SỐ CỤ THỂ NÀY?', 1)

    add_image_para(doc, 'diagram_09_hyperparams.png', Inches(6.5))
    p = doc.add_paragraph('Hình 7: Giải thích 4 hyperparameters chính')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    add_heading_custom(doc, '7.1 Bảng tổng hợp hyperparameters', 2)

    add_body(doc, 'n_lags = 12: 12 x 5 phút = 60 phút lookback - đủ capture intra-hour patterns')
    add_body(doc, 'n_forecasts = 6: 6 x 5 phút = 30 phút horizon - phù hợp controller planning')
    add_body(doc, 'ar_layers = [32, 32]: Sweet spot cho data size ~20-36 điểm')
    add_body(doc, 'quantiles = [0.1, 0.9]: 80% interval - vừa đủ rộng, vừa actionable')

    add_heading_custom(doc, '7.2 So sánh chi tiết n_lags', 2)

    add_body(doc, 'n_lags = số timestep QUÁ KHỨ được dùng để dự báo')

    headers_nlags = ['Giá trị', 'Thời gian lookback', 'Ưu điểm', 'Nhược điểm', 'Đánh giá']
    rows_nlags = [
        ['n_lags=6', '30 phút', 'Ít data, nhanh', 'Có thể miss pattern dài', 'Ngắn quá'],
        ['n_lags=12', '60 phút', 'Đủ capture intra-hour', 'Balance tốt', 'TỐI ƯU'],
        ['n_lags=24', '120 phút', 'Rất nhiều context', 'Overfit với data nhỏ', 'Dư thừa'],
        ['n_lags=48', '240 phút', 'Very long context', 'Overfit nặng, chậm', 'Không phù hợp'],
    ]
    add_table(doc, headers_nlags, rows_nlags)

    add_spacer(doc, 0.2)
    add_body(doc, 'Kết luận: n_lags=12 là tối ưu vì 60 phút đủ để bắt các pattern trong giờ (intra-hour patterns) mà không gây overfit.')

    add_heading_custom(doc, '7.3 So sánh chi tiết ar_layers', 2)

    add_body(doc, 'ar_layers = cấu trúc Neural Network bên trong AR-Net')

    headers_arlayers = ['Cấu trúc', 'Số params ước tính', 'Phù hợp với', 'Overfit risk', 'Đánh giá']
    rows_arlayers = [
        ['ar_layers=[16]', '~500 params', 'Data rất ít (<15 điểm)', 'Thấp', 'Có thể underfit'],
        ['ar_layers=[32]', '~1,000 params', 'Data ít (15-25 điểm)', 'Trung bình', 'Khá tốt'],
        ['ar_layers=[32, 32]', '~2,100 params', 'Data vừa (25-50 điểm)', 'Thấp', 'TỐI ƯU'],
        ['ar_layers=[64, 64]', '~8,000 params', 'Data nhiều (>100 điểm)', 'Cao', 'Quá phức tạp'],
        ['ar_layers=[64, 64, 64]', '~20,000+ params', 'Data rất nhiều (>500 điểm)', 'Rất cao', 'Overkill'],
    ]
    add_table(doc, headers_arlayers, rows_arlayers)

    add_spacer(doc, 0.2)
    add_body(doc, 'Kết luận: ar_layers=[32, 32] là sweet spot cho dataset ~30 điểm. 2 layers đủ deep để học non-linear patterns mà không overfit.')

    add_heading_custom(doc, '7.4 So sánh chi tiết quantiles', 2)

    add_body(doc, 'quantiles = khoảng dao động confidence interval')

    headers_quant = ['Quantiles', 'Coverage', 'Độ rộng', 'Phù hợp cho', 'Đánh giá']
    rows_quant = [
        ['[0.1, 0.9]', '80%', 'Rộng vừa', 'Controller planning', 'TỐI ƯU'],
        ['[0.2, 0.8]', '60%', 'Hẹp', 'High confidence cần', 'Hơi hẹp'],
        ['[0.05, 0.95]', '90%', 'Rất rộng', 'Conservative decision', 'Quá rộng'],
        ['[0.25, 0.75]', '50%', 'Rất hẹp', 'Chỉ show median', 'Không đủ useful'],
        ['[0.1, 0.5, 0.9]', '80% (3 outputs)', 'Bình thường', 'Khi cần 3 values', 'Tốt nhưng phức tạp hơn'],
    ]
    add_table(doc, headers_quant, rows_quant)

    add_spacer(doc, 0.2)
    add_body(doc, 'Kết luận: quantiles=[0.1, 0.9] với 80% interval là sweet spot - đủ rộng để coverage tốt, đủ hẹp để actionable cho controller.')

    add_heading_custom(doc, '7.5 So sánh n_forecasts', 2)

    add_body(doc, 'n_forecasts = số bước dự báo tương lai')

    headers_forecasts = ['Giá trị', 'Thời gian horizon', 'Phù hợp cho', 'Ưu điểm', 'Nhược điểm', 'Đánh giá']
    rows_forecasts = [
        ['n_forecasts=3', '15 phút', 'Very short-term', 'Ít uncertainty', 'Chỉ kịp phản ứng', 'Ngắn quá'],
        ['n_forecasts=6', '30 phút', 'Controller planning', 'Balance tốt', 'Vừa đủ', 'TỐI ƯU'],
        ['n_forecasts=12', '60 phút', 'Strategic planning', 'Dài hơn', 'Uncertainty cao', 'Hơi dài'],
        ['n_forecasts=18', '90 phút', 'Very long horizon', 'Rất dài', 'Uncertainty rất cao', 'Không phù hợp'],
    ]
    add_table(doc, headers_forecasts, rows_forecasts)

    add_spacer(doc, 0.2)
    add_body(doc, 'Kết luận: n_forecasts=6 (30 phút) là phù hợp nhất vì controller thường plan theo 30-phút blocks.')

    add_spacer(doc, 0.3)

    # ==================== PHAN 8: COMPARISON ====================
    add_heading_custom(doc, 'PHẦN 8: SO SÁNH CÁC PHƯƠNG PHÁP DỰ BÁO', 1)

    add_image_para(doc, 'diagram_07_comparison.png', Inches(7))
    p = doc.add_paragraph('Hình 8: So sánh NeuralProphet với ARIMA, LSTM, Transformer')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.italic = True

    add_spacer(doc, 0.2)

    add_heading_custom(doc, '8.1 Bảng so sánh các phương pháp', 2)

    headers = ['Tiêu chí', 'NeuralProphet', 'ARIMA', 'LSTM', 'Transformer']
    rows = [
        ['Độ phức tạp', 'Trung bình', 'Thấp', 'Cao', 'Rất cao'],
        ['Data cần thiết', '~20 điểm', '~100 điểm', '~1000 điểm', '~5000 điểm'],
        ['Non-linear', 'Có', 'Không', 'Có', 'Có'],
        ['Uncertainty', 'Có (Quantiles)', 'Không', 'Không', 'Có'],
        ['Đa biến (Weather)', 'Dễ dàng', 'Phức tạp', 'Dễ dàng', 'Dễ dàng'],
        ['Interpretable', 'Cao', 'Cao', 'Thấp', 'Rất thấp'],
        ['Training Speed', 'Nhanh', 'Nhanh', 'Chậm', 'Rất chậm'],
        ['Small data (<50)', 'Rất phù hợp', 'Phù hợp', 'Không phù hợp', 'Không phù hợp'],
    ]
    add_table(doc, headers, rows)

    add_spacer(doc, 0.3)
    add_heading_custom(doc, '8.2 Tại sao chọn NeuralProphet?', 2)

    add_bullet(doc, 'THIẾT KẾ CHO DATA NHỎ (~20+ điểm là đủ)')
    add_bullet(doc, 'AR-NET: Kết hợp autoregressive + neural network - tốt nhất cả 2 thế giới')
    add_bullet(doc, 'Native support cho exogenous variables (weather, buffer)')
    add_bullet(doc, 'Uncertainty estimation với quantile regression')
    add_bullet(doc, 'Interpretable - dễ giải thích với thầy')

    add_spacer(doc, 0.3)

    # ==================== PHAN 9: WHY NOT OTHER METHODS ====================
    add_heading_custom(doc, 'PHẦN 9: TẠI SAO KHÔNG DÙNG PHƯƠNG PHÁP KHÁC?', 1)

    add_heading_custom(doc, '9.1 Tại sao không dùng ARIMA?', 2)

    add_body(doc, 'VẤN ĐỀ 1: ARIMA chỉ học được QUAN HỆ TUYẾN TÍNH')
    add_bullet(doc, 'ARIMA giả định: y = a1*x1 + a2*x2 + ... + constant')
    add_bullet(doc, '-> Luôn là đường THẲNG')
    add_bullet(doc, 'NHƯNG thực tế: Giờ cao điểm 5->10 máy bay (tăng GẤP ĐÔI), giờ thường 5->6 (tăng 20%)')
    add_bullet(doc, '-> Đây là NON-LINEAR!')

    add_body(doc, 'VẤN ĐỀ 2: ARIMA khó dùng với nhiều biến')
    add_bullet(doc, 'Muốn dùng weather + buffer trong ARIMA: phải dùng ARIMAX (phức tạp)')
    add_bullet(doc, 'NeuralProphet: add_future_regressor() và add_lagged_regressor() - Đơn giản!')

    add_body(doc, 'VẤN ĐỀ 3: ARIMA không có Uncertainty estimation')
    add_bullet(doc, 'ARIMA: "Có 6 máy bay" -> Không biết khoảng dao động')
    add_bullet(doc, 'NeuralProphet: "Có 6 máy bay (4-8)" -> Có khoảng dao động 80% confidence')

    add_heading_custom(doc, '9.2 Tại sao không dùng LSTM?', 2)

    add_body(doc, 'VẤN ĐỀ CƠ BẢN: LSTM CẦN RẤT NHIỀU DATA')
    add_bullet(doc, 'LSTM được thiết kế cho: Stock price (10 năm data), NLP (hàng triệu từ), Video (hàng triệu frames)')
    add_bullet(doc, 'Dataset của chúng ta: ~20-36 điểm sau vài giờ thu thập')
    add_bullet(doc, 'Sự khác biệt: 1:100,000+')

    add_body(doc, 'NẾU DÙNG LSTM VỚI DATA NHỎ SẼ XẢY RA:')
    add_bullet(doc, 'OVERFITTING: Model "nhớ" toàn bộ training data thay vì học pattern')
    add_bullet(doc, 'Training error: RẤT THẤP nhưng Test error: RẤT CAO')
    add_bullet(doc, 'Model KHÔNG CÓ GIÁ TRỊ thực tế')

    add_heading_custom(doc, '9.3 Tại sao không dùng Prophet (Facebook)?', 2)

    add_body(doc, 'Prophet = Phương pháp gốc, NeuralProphet = Phiên bản CẢI TIẾN')
    add_bullet(doc, 'Prophet: KHÔNG CÓ autoregressive component!')
    add_bullet(doc, 'NeuralProphet: CÓ AR-Net!')
    add_bullet(doc, 'Ví dụ: 10:00 có 5 máy bay, 10:05 có 7 máy bay')
    add_bullet(doc, 'Prophet: Không dùng thông tin này để predict 10:10')
    add_bullet(doc, 'NeuralProphet: Dùng 5,7 để predict 10:10 -> tốt hơn')

    add_spacer(doc, 0.3)

    # ==================== PHAN 10: GLOSSARY ====================
    add_heading_custom(doc, 'PHẦN 10: CÁC THUẬT NGỮ VÀ GIẢI THÍCH', 1)

    glossary = [
        ('AR-Net', 'Autoregressive Neural Network - Neural network dùng quá khứ để dự báo tương lai'),
        ('Autoregressive', '"Tự hồi quy" - dùng giá trị trước đó để dự báo giá trị tiếp theo'),
        ('Backpropagation', 'Thuật toán để neural network "học" - tính sai số và điều chỉnh weights'),
        ('Epoch', 'Một lần "học" toàn bộ dataset'),
        ('Exogenous Variable', 'Biến ngoại sinh - yếu tố bên ngoài ảnh hưởng đến kết quả'),
        ('Forecast Horizon', 'Độ dài thời gian dự báo (30 phút trong project này)'),
        ('Lagged Regressor', 'Biến trễ - giá trị QUÁ KHỨ của một biến được dùng làm input'),
        ('Neuron', 'Đơn vị tính toán cơ bản trong neural network'),
        ('Non-linear', 'Không theo đường thẳng - thực tế phức tạp hơn'),
        ('Overfitting', 'Model "học vẹt" - nhớ training data thay vì học pattern'),
        ('Quantile', 'Phân vị - giúp ước lượng uncertainty'),
        ('Seasonality', 'Tính chu kỳ - pattern LẶP LẠI theo ngày/tuần/năm'),
        ('Underfitting', 'Model quá đơn giản - không học được pattern'),
        ('Weights', 'Trọng số - độ quan trọng của mỗi kết nối trong neural network'),
    ]

    for term, definition in glossary:
        p = doc.add_paragraph()
        run = p.add_run(f'{term}: ')
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(definition)
        run.font.size = Pt(11)

    add_spacer(doc, 0.3)

    # ==================== PHAN 11: CONCLUSION ====================
    add_heading_custom(doc, 'PHẦN 11: KẾT LUẬN VÀ CÂU HỎI THƯỜNG GẶP', 1)

    add_heading_custom(doc, '11.1 Tóm tắt kiến trúc model', 2)

    add_body(doc, 'NeuralProphet = Sự KẾT HỢP hoàn hảo giữa:')
    add_bullet(doc, 'Độ ĐƠN GIẢN của ARIMA (interpretable, fast)')
    add_bullet(doc, 'Độ MẠNH MẼ của Neural Network (non-linear)')
    add_bullet(doc, 'Độ LINH HOẠT cho small data (thiết kế riêng)')

    add_heading_custom(doc, '11.2 Cấu hình hiện tại có phải là tối ưu?', 2)

    add_body(doc, 'Trả lời: GẦN NHƯ TỐI ƯU CHO HIỆN TẠI')

    add_bullet(doc, 'n_lags=12: Phù hợp (60 phút lookback)')
    add_bullet(doc, 'n_forecasts=6: Phù hợp (30 phút horizon)')
    add_bullet(doc, 'ar_layers=[32,32]: Sweet spot cho data size')
    add_bullet(doc, 'quantiles=[0.1,0.9]: 80% interval phù hợp')

    add_heading_custom(doc, '11.3 Câu hỏi thường gặp từ thầy', 2)

    faqs = [
        ('Q1: Tại sao dùng NeuralProphet mà không dùng ARIMA?',
         'ARIMA chỉ học được quan hệ TUYẾN TÍNH, trong khi air traffic có pattern NON-LINEAR.'),
        ('Q2: Tại sao n_lags=12 mà không phải 6 hay 24?',
         '12 lags = 60 phút = đủ capture intra-hour patterns mà không overfit.'),
        ('Q3: Tại sao ar_layers=[32,32] mà không phải [16] hay [64,64]?',
         '[16] có thể underfit, [64,64] có thể overfit. [32,32] là sweet spot.'),
        ('Q4: Sao không dùng LSTM? Đó là state-of-the-art mà?',
         'LSTM cần ~1000+ điểm data. Dataset chỉ có ~30 điểm -> LSTM sẽ OVERFIT nặng.'),
        ('Q5: Khoảng dao động 80% lấy ở đâu?',
         '80% interval là sweet spot: đủ rộng để coverage tốt, đủ hẹp để actionable.'),
        ('Q6: Làm sao model biết weather ảnh hưởng thế nào?',
         'Weather được thêm như regressors. Trong training, model học coefficients.'),
        ('Q7: Buffer count là gì?',
         'Buffer zones là các vùng xung quanh sector. buffer_n/s/e/w = số máy bay đang tiến vào.'),
        ('Q8: Model có thể predict khi nào DỰ ĐOÁN SAI không?',
         'Khoảng dao động chính là để handle điều này. Controller được training để hiểu uncertainty.'),
    ]

    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(q)
        run.bold = True
        run.font.size = Pt(11)
        p = doc.add_paragraph(a)
        for run in p.runs:
            run.font.size = Pt(11)
        add_spacer(doc, 0.1)

    # ==================== APPENDIX ====================
    add_heading_custom(doc, 'PHỤ LỤC: CÔNG THỨC TOÁN HỌC', 1)

    add_heading_custom(doc, 'Observation Model', 2)

    add_body(doc, 'y(t) = f(y(t-1), y(t-2), ..., y(t-12);')
    add_body(doc, '           x_weather(t);')
    add_body(doc, '           z_buffer(t-1), ..., z_buffer(t-12)) + e(t)')

    add_body(doc, 'Trong đó:')
    add_bullet(doc, 'y(t): Số máy bay tại thời điểm t (target variable)')
    add_bullet(doc, 'y(t-k): lagged target values (autoregressive)')
    add_bullet(doc, 'x_weather: [precipitation, wind_speed, cloud_cover, visibility]')
    add_bullet(doc, 'z_buffer: [buffer_n, buffer_s, buffer_e, buffer_w]')
    add_bullet(doc, 'e(t): noise term')
    add_bullet(doc, 'f(): neural network function (AR-Net)')

    add_heading_custom(doc, 'Neural Network Architecture', 2)

    add_body(doc, 'h1 = ReLU(W1 * [y(t-12:t), x(t), z(t-12:t)] + b1)')
    add_body(doc, 'h2 = ReLU(W2 * h1 + b2)')
    add_body(doc, 'y_hat(t+1:t+6) = W3 * h2 + b3')

    add_heading_custom(doc, 'Loss Function (Quantile Regression)', 2)

    add_body(doc, 'L(theta) = Sum(q in {0.1,0.5,0.9}) Sum rho_q(y_i - y_hat_q,i)')
    add_body(doc, 'trong đó rho_q(u) = u*(q - I(u<0)) [pinball loss]')

    add_heading_custom(doc, 'Prediction Interval', 2)

    add_body(doc, 'P(y in [y_hat_10, y_hat_90]) ~ 0.80')
    add_body(doc, '80% confidence rằng giá trị thực nằm trong khoảng [y_10, y_90]')

    # Save
    doc.save(SPEC_DOCX)
    print(f"Document saved to: {SPEC_DOCX}")

if __name__ == '__main__':
    create_document()
